import streamlit as st
from openai import OpenAI
import os
from docx import Document
from io import BytesIO

# ------------------------
# Config
# ------------------------
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

if not OPENAI_API_KEY:
    st.error("❌ OPENAI_API_KEY not set in environment variables.")
    st.stop()

client = OpenAI(api_key=OPENAI_API_KEY)

# ------------------------
# Helpers
# ------------------------
def read_docx(file):
    doc = Document(file)
    text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    return text

def create_docx_clean(text):
    doc = Document()
    lines = text.split("\n")

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Section headings
        if line.upper() in [
            "SUMMARY", "SKILLS", "INDUSTRIAL WORKING EXPERIENCE",
            "CORPORATE TRAINING EXPERIENCE", "RETAIL TRAINING EXPERIENCE",
            "MATERIAL EXPERIENCE", "CERTIFICATIONS", "EDUCATION", "PERSONAL DETAILS"
        ]:
            para = doc.add_paragraph(line)
            para.style = "Heading 1"
        elif line.startswith(("-", "•")):
            clean = line.lstrip("-• ").strip()
            para = doc.add_paragraph(clean)
            para.style = "List Bullet"
        else:
            para = doc.add_paragraph(line)
            para.style = "Normal"

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def ask_chatgpt(requirements, profile_text, user_extra, model, temperature):
    base_prompt = f"""
You are an expert CV/profile rewriting assistant.

TASK:
- Rewrite the trainer profile using the given requirements.
- Keep the document structure professional and ATS-friendly.
- Section order:
  1. Title(Use Heading 1 include Name - Mujahed H.) & Sub title(use Heading 2 suggest best title like: Cloud DevOps Trainer, AI Trainer) 
  2. Summary(This should be match what requirements weshared with you)
  3. Skills (with categories AND representative technology keywords, e.g. Big Data (Spark, Kafka, Hadoop) or DevOps (Jenkins, Ansible, Terraform, Github))
  4. Industrial Working Experience (per client: Client, Focus, Technologies (categories + key tools), Achievements)
  5. Corporate Training Experience
  6. Retail Training Experience
  7. Material Experience
  8. Certifications
  9. Education
  10. Personal Details
- Replace long raw tool dumps with concise categories, but keep keyword coverage for ATS.
- Output plain text only (no Markdown # or **).

REQUIREMENTS:
{requirements}

PROFILE DATA:
{profile_text}

EXTRA USER INSTRUCTIONS:
{user_extra}
"""

    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": "You are a professional trainer profile generator."},
            {"role": "user", "content": base_prompt}
        ],
        temperature=temperature,
    )

    return response.choices[0].message.content

# ------------------------
# Streamlit UI
# ------------------------
st.set_page_config(page_title="ATS-Friendly Trainer Profile Generator", layout="wide")

st.title("📄 ATS-Friendly Trainer Profile Generator (ChatGPT API)")

st.sidebar.header("⚙️ Settings")
model = st.sidebar.selectbox(
    "Choose a model:",
    [
        "gpt-4o-mini",
        "gpt-4o",
        "gpt-4-turbo",
        "gpt-3.5-turbo"
    ],
    index=0
)
temperature = st.sidebar.slider("Creativity (temperature)", 0.0, 1.5, 0.7, 0.1)

req_file = st.file_uploader("📂 Upload Requirements Document (.docx)", type=["docx"])
profile_file = st.file_uploader("📂 Upload Profile Document (.docx)", type=["docx"])

user_extra = st.text_area(
    "✍️ Extra Instructions (optional)",
    placeholder="E.g., Emphasize Big Data projects, shorten Retail Training, make summary stronger..."
)

if st.button("🚀 Generate Profile"):
    if req_file and profile_file:
        with st.spinner("Generating ATS-friendly profile..."):
            requirements_text = read_docx(req_file)
            profile_text = read_docx(profile_file)

            rewritten_text = ask_chatgpt(requirements_text, profile_text, user_extra, model, temperature)

            final_doc = create_docx_clean(rewritten_text)

            st.success("✅ ATS-friendly trainer profile generated!")
            st.download_button(
                label="📥 Download Final Profile",
                data=final_doc,
                file_name="trainer_profile_final.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            st.subheader("Preview")
            st.text_area("AI Output (preview only):", rewritten_text, height=400)
    else:
        st.error("⚠️ Please upload both documents before generating.")
